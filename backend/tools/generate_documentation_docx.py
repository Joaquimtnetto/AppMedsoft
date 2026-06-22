#!/usr/bin/env python3
"""
Gera DOCUMENTATION.docx a partir de DOCUMENTATION.md e DOCUMENTATION_DIAGRAMS.md.

Comportamento:
- Concatena os conteúdos dos dois arquivos em um documento .docx legível.
- Para blocos Mermaid (```mermaid ... ```), tenta gerar PNG usando `mmdc` (mermaid-cli) se disponível e embutir as imagens.

Requisitos para geração de imagens (opcional):
- Instalar mermaid-cli: `npm install -g @mermaid-js/mermaid-cli`
- Instalar python-docx: `pip install python-docx`

Uso:
  python tools/generate_documentation_docx.py

Saída:
  backend/DOCUMENTATION.docx
"""
import os
import re
import subprocess
from pathlib import Path
try:
    from docx import Document
    from docx.shared import Inches
except Exception:
    print('Erro: instale python-docx com `pip install python-docx`')
    raise

ROOT = Path(__file__).resolve().parents[1]
DOC_MD = ROOT / 'DOCUMENTATION.md'
DIAG_MD = ROOT / 'DOCUMENTATION_DIAGRAMS.md'
OUT_DOCX = ROOT / 'DOCUMENTATION.docx'
TEMP_DIR = ROOT / 'tools' / 'docx_temp'
TEMP_DIR.mkdir(parents=True, exist_ok=True)

MERMAID_CLI = shutil_which = None
def find_mmdc():
    # procura por mmdc no PATH
    from shutil import which
    return which('mmdc') or which('mermaid-cli')

def render_mermaid(code, out_png_path):
    mmdc = find_mmdc()
    if not mmdc:
        return False
    # escreve código temporário
    tmp_md = TEMP_DIR / 'diag_tmp.mmd'
    tmp_md.write_text(code, encoding='utf-8')
    try:
        subprocess.check_call([mmdc, '-i', str(tmp_md), '-o', str(out_png_path)])
        return True
    except Exception:
        return False

def add_markdown_to_doc(doc, md_path):
    text = md_path.read_text(encoding='utf-8')
    # separar por linhas para preservar
    lines = text.splitlines()
    in_code = False
    code_lang = None
    code_buf = []
    for line in lines:
        m_start = re.match(r'^```(\w*)', line)
        if m_start:
            in_code = True
            code_lang = m_start.group(1)
            code_buf = []
            continue
        if in_code and line.strip() == '```':
            # fim do bloco
            in_code = False
            if code_lang == 'mermaid':
                # tentar renderizar
                png_name = 'diagram_{}.png'.format(abs(hash('\n'.join(code_buf))) % (10**8))
                out_png = TEMP_DIR / png_name
                rendered = render_mermaid('\n'.join(code_buf), out_png)
                if rendered and out_png.exists():
                    doc.add_picture(str(out_png), width=Inches(6))
                else:
                    doc.add_paragraph('Mermaid diagram (não renderizado):')
                    doc.add_paragraph('\n'.join(code_buf))
            else:
                # adicionar bloco de código como parágrafo pré-formatado
                doc.add_paragraph('Código (%s):' % (code_lang or ''))
                for c in code_buf:
                    p = doc.add_paragraph(c)
                    p.style = 'Intense Quote'
            code_lang = None
            code_buf = []
            continue
        if in_code:
            code_buf.append(line)
        else:
            # tratar headings markdown
            h = re.match(r'^(#{1,6})\s+(.*)', line)
            if h:
                lvl = len(h.group(1))
                text_h = h.group(2)
                if lvl == 1:
                    doc.add_heading(text_h, level=1)
                elif lvl == 2:
                    doc.add_heading(text_h, level=2)
                else:
                    doc.add_heading(text_h, level=3)
            else:
                if line.strip() == '---':
                    # separador
                    doc.add_page_break()
                else:
                    doc.add_paragraph(line)

def main():
    doc = Document()
    doc.core_properties.title = 'DOCUMENTATION'
    if DOC_MD.exists():
        add_markdown_to_doc(doc, DOC_MD)
    if DIAG_MD.exists():
        add_markdown_to_doc(doc, DIAG_MD)
    doc.save(str(OUT_DOCX))
    print('Gerado:', OUT_DOCX)

if __name__ == '__main__':
    main()
